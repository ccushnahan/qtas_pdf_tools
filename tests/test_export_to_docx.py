"""Tests export to docx module"""
import pytest
from qtas_pdf.export_to_docx import *
import os
import docx


def test_find_tables_in_pdf_invalid_filename():
    """Test find tables."""
    file_name = "this_doesnt_exist.pdf"
    with pytest.raises(FileNotFoundError) as e_info:
        find_tables_in_pdf(file_name)

def test_find_tables_in_pdf_invalid_filetype():
    """Test find tables."""
    file_name = "requirements.txt"
    with pytest.raises(NotImplementedError) as e:
        find_tables_in_pdf(file_name)
    
def test_find_tables_no_tables():
    """Test find tables."""
    file_name = "tests/test_data/blank.pdf"
    data_frame = find_tables_in_pdf(file_name)
    assert data_frame == None

def test_find_tables_single_table_row_count():
    """Test find tables."""
    file_name = "tests/test_data/unit_1.pdf"
    data_frame = find_tables_in_pdf(file_name)
    assert len(data_frame.index) == 8

def test_find_tables_correct_return():
    """Test find tables."""
    file_name = "tests/test_data/unit_1.pdf"
    data_frame = find_tables_in_pdf(file_name)
    assert type(data_frame) == pandas.DataFrame

def test_process_pdf():
    """Test process pdf"""

def test_clean_and_structure_tables_data_type():
    """Test clean and structure tables"""
    file_name = "tests/test_data/unit_1.pdf"
    data_frame = find_tables_in_pdf(file_name)
    clean_data = clean_and_structure_tables(data_frame)
    assert type(clean_data) == list

def test_clean_and_structure_tables_number_cleaned():
    """Test clean and structure tables"""
    file_name = "tests/test_data/unit_1.pdf"
    data_frame = find_tables_in_pdf(file_name)
    clean_data = clean_and_structure_tables(data_frame)
    print(clean_data)
    assert len(clean_data) == 2

def test_clean_and_structure_tables():
    """Test clean and structure tables"""
    with pytest.raises(Exception) as e:
        file_name = "tests/test_data/blank.pdf"
        data_frame = find_tables_in_pdf(file_name)
        clean_data = clean_and_structure_tables(data_frame)

def test_insert_tables_into_docx_invalid_data():
    """Test insert tables into docx"""
    with pytest.raises(Exception) as e:
        insert_tables_into_docx(None, "")

def test_insert_tables_into_docx_valid_file():
    """Test insert tables into docx"""
    file_name = "tests/test_data/unit_1.pdf"
    data_frame = find_tables_in_pdf(file_name)
    clean_data = clean_and_structure_tables(data_frame)
    insert_tables_into_docx(clean_data, "test")
    assert os.path.exists("test.docx")
    
def test_set_signature_section_no_doc():
    """Test set signature section"""
    with pytest.raises(Exception) as e:
        set_signature_section(None)

def test_set_signature_section_valid_doc():
    """Test set signature section"""
    doc = docx.Document()
    doc = set_signature_section(doc)
    assert len(doc.paragraphs) == 5
    assert doc.paragraphs[2].text == "Learner Name:\t\t"
    
def test_set_header_no_doc():
    """Test set header"""
    with pytest.raises(Exception) as e:
        set_header(None)

def test_set_header_valid_doc():
    """Test set header"""
    doc = docx.Document()
    doc = set_header(doc)
    assert len(doc.sections[0].header.paragraphs) == 2
    assert doc.sections[0].header.paragraphs[0].text == "QTAS"

def test_main_method():
    main()
    assert os.path.exists("unit_25.docx")