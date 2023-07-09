"""Takes pdf file, extracts, transforms then exports tables to docx.

Export to docx opens pdf files and identifies tables within files. It
extracts those tables from the file, rearranges the table format and 
then inserts the reformatted table into a docx document with appropriate
header for unit name.

Typical Usage:

    python3 export_to_pdf.py path_to_pdf

    OR

    for file in os.listdir(split_files_location):
        file_path = os.path.join(split_files_location, file)
        export_to_docx.process_pdf(file_path)
"""

import re
import sys

import camelot
import docx
import pandas
import PyPDF2
import tabula
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger


def process_pdf(filename):
    """Process the pdf to output doc.

    Args:
        filename (str): filename
    """
    logger.info("Finding Tables...")
    tables = find_tables_in_pdf(filename)
    logger.info("Cleaning Tables")
    clean_tables = clean_and_structure_tables(tables)
    logger.info("Inserting Tables...")
    insert_tables_into_docx(clean_tables, filename)


def find_tables_in_pdf(filename: str) -> pandas.DataFrame:
    """Find the tables in the pdf.

    Take pdf file path and open it. Run camelot.read_pdf
    over the document and extract the tables.

    Args:
        filename (str): Filename/path string

    Returns:
        joined_table (pandas.DataFrame): Tables concat together

    """
    logger.debug("Open pdf and read tables")
    logger.debug("Opening pdf...")
    logger.debug("Reading tables...")
    tables = camelot.read_pdf(filename, pages="1-end", strip_text="\n")
    logger.debug(f"Tables Found: {len(tables)}")
    joined_tables = None
    for table in tables:
        if joined_tables is None:
            joined_tables = table.df
        else:
            joined_tables = pandas.concat([joined_tables, table.df])
    return joined_tables

def clean_and_structure_tables(table: pandas.DataFrame) -> pandas.DataFrame:
    """Clean and restructure tables.

    Extract table data from tables list. Transform into
    pandas.Dataframes and then clean and manipulate the
    data to return a new list of tables in the structure
    that they will be inserted into the Word (.docx)
    document.

    Args:
        tables (list):  List of tables found from camelot.read_pdf

    Returns:
        clean_tables (list): List of pandas.Dataframe objects
    """
    sometables = []
    titles = []

    logger.debug("Iterate through tables to clean and structure table data.")

    c = table
    logger.debug("Cleaning table")
    c = c.drop([0])
    c = c[c[2] != ""]
    c[1] = c[1].map(lambda x: re.sub(r"[^A-Za-z0-9.: ]", "", str(x)).strip())
    c[1] = c[1].map(lambda x: None if x == "nan" or x == "" else x)
    c[2] = c[2].map(lambda x: re.sub(r"[^A-Za-z0-9.: ]", "", str(x)).strip())
    c[3] = c[3].map(lambda x: re.sub(r"[^A-Za-z0-9.: ]", "", str(x)).strip())
    c["titles"] = c[0].astype(str) + ". " + c[1]
    titles = [*titles, *c["titles"].dropna().to_list()]
    for i in range(len(titles)):
        if not titles[i][0].isdigit():
            titles[i - 1] += titles[i]
    titles = [*filter(lambda x: x[0].isdigit(), titles)]
    c["outcomes"] = c[2].astype(str) + " " + c[3]
    columns = c[[2, 3, 4, "outcomes"]]
    last_section = c.iloc[-1]
    last_section_num = str(last_section[2]).split(".")[0]
    for i in range(1, int(last_section_num) + 1):
        sub_sec = columns[columns[2] != ""]
        sub_sec = columns[columns[2].astype(float) > i]
        sub_sec = sub_sec[sub_sec[2].astype(float) < (i + 1)]
        sub_sec = sub_sec[["outcomes", 4]]
        title = ""
        logger.debug("Finding table title")
        try:
            title = titles[i - 1]
        except:
            title = ""
        logger.debug(f"Table Title: {title}")

        sub_sec.rename(columns={"outcomes": f"{title}", 4: "Evidence"}, inplace=True)
        if len(sub_sec.index) >= 1:
            sometables.append(sub_sec)
    return sometables


def insert_tables_into_docx(tables: list, filename: str) -> None:
    """Insert the tables into docx file.
    
    Takes list of tables and a file name. Creates new docx
    Document and iterates over the list of tables to insert 
    into the document. Saves the document under the filename.

    Args:
        tables (list):      List of pandas.Dataframes
        filename (str):     Name for output file 
    """
    doc = docx.Document()
    doc = set_header(doc)
    heading = " ".join(filename.split("/")[-1].split("_")).split(".")[0].title()
    doc.add_heading(heading, 0)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    sometables = tables
    fileparts = filename.split("/")
    filename = fileparts[-1]
    fileparts = filename.split(".")
    doc_name = fileparts[0] + ".docx"

    for sometable in sometables:
        logger.debug("Shaping word table based on df")
        t = doc.add_table(sometable.shape[0] + 3, sometable.shape[1])
        t.allow_autofit = True
        t.autofit = True
        t.style = "Table Grid"

        logger.debug("Inserting data into word table")
        for j in range(sometable.shape[-1]):
            t.rows[0].height = 800000
            t.rows[1].height = 300000
            if j == 0:
                t.cell(0, j).text = "\n" + sometable.columns[j]
                t.cell(0, j).merge(t.cell(0, j + 1))
                t.cell(1, j).text = "\nCriteria"
                t.cell(0, j).paragraphs[0].runs[0].font.bold = True
                t.cell(1, j).paragraphs[0].runs[0].font.bold = True
            if j == 1:
                t.cell(1, j).text = "\nEvidence"
                t.cell(1, j).paragraphs[0].runs[0].font.bold = True

            for i in range(sometable.shape[0]):
                t.cell(i + 2, 0).text = "\n" + str(sometable.values[i, 0]) + "\n"
                t.cell(2, 1).merge(t.cell(i + 3, 1))
            if i < (len(sometable.index) - 3):
                t.cell(i + 2, 0).merge(t.cell(i + 2, 1))
        t.columns[0].width = 4560065
        t.columns[1].width = 1560065

    doc = set_signature_section(doc)
    logger.debug("Saving Word data")
    doc.save(doc_name)


def set_signature_section(doc: docx.Document) -> docx.Document:
    """Adds a signature section.
    
    Adds a signature section to end of tables to allow for
    learner and verifier sign off of courses.

    Args:
        doc (docx.Document):    Document to which the section is
                                added.
    """
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("Learner Name:\t\t")
    doc.add_paragraph("Learner Signature:\t\t\t\t\t\tAssessor:\t\t\t")
    doc.add_paragraph("Internal Verifier:\t\t\t\t\t\tDate:\t\t\t")

    return doc


def set_header(doc: docx.Document) -> docx.Document:
    """Set and insert document header.
    
    Creates and formats a header for the document based
    on QTAS requirements.

    Args:
        doc (docx.Document):    Document to which the header
                                is added.
    """
    header = doc.sections[0].header
    header.add_paragraph()
    header_one = header.paragraphs[0]
    header_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # header_one.runs[0].font.color =
    header_one.text = "QTAS"
    header_one.style.font.size = docx.shared.Pt(18)
    header_two = header.paragraphs[1]
    header_two.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_two.text = "Quality Training and Assessment Services\n"
    header_two.style.font.size = docx.shared.Pt(11)
    return doc


def main():
    """The Main method."""
    try:
        pdf_file = sys.argv[1]
    except:
        pdf_file = "data/units/unit_25.pdf"
    process_pdf(pdf_file)


if __name__ == "__main__":
    main()
